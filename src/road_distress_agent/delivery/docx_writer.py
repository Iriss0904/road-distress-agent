"""Render a formal inspection report to a .docx document."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from road_distress_agent.delivery.docx_format import (
    TITLE_TEXT,
    add_formal_paragraph,
    add_heading,
    add_labeled_paragraph,
    add_subheading,
    add_text,
    add_value_text,
    configure_formal_document,
)
from road_distress_agent.delivery.report_labels import report_label


@dataclass(frozen=True)
class ReportDocumentContext:
    doc: Any
    project: dict[str, Any]
    report_meta: dict[str, Any]
    locale: str


def write_report_document(
    *,
    project: dict[str, Any],
    report_meta: dict[str, Any],
    overall_summary: str,
    defect_sections: list[dict[str, Any]],
    cost_summary: dict[str, Any] | None,
    path: Path,
    locale: str = "zh-CN",
) -> Path:
    doc = Document()
    configure_formal_document(doc)
    context = ReportDocumentContext(
        doc=doc,
        project=project,
        report_meta=report_meta,
        locale=locale,
    )
    _write_title(context)
    _write_overview(context, len(defect_sections))
    _write_results(context, overall_summary, defect_sections)
    _write_recommendations(context, defect_sections, cost_summary)
    _write_attachments(context)
    doc.save(path)
    return path


def _write_title(context: ReportDocumentContext) -> None:
    paragraph = add_formal_paragraph(context.doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_value_text(
        paragraph,
        context.report_meta.get("unit_name"),
        _label("unit_name", context.locale),
        style=TITLE_TEXT,
        locale=context.locale,
    )
    add_text(paragraph, f" {_label('title', context.locale)}", style=TITLE_TEXT)
    paragraph = add_formal_paragraph(context.doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(
        paragraph,
        f"{_label('report_date', context.locale)}: {_report_date(context.locale)}",
    )


def _write_overview(context: ReportDocumentContext, defect_count: int) -> None:
    add_heading(context.doc, _label("overview", context.locale))
    _labeled(context, _label("task", context.locale), context.project.get("name"))
    _labeled(
        context,
        _label("inspection_time", context.locale),
        context.report_meta.get("inspection_time"),
    )
    _labeled(context, _label("segment", context.locale), context.report_meta.get("segment"))
    _labeled(
        context,
        _label("inspection_method", context.locale),
        context.report_meta.get("inspection_method"),
    )
    _labeled(
        context,
        _label("participants", context.locale),
        context.report_meta.get("participants"),
    )
    _labeled(context, _label("crew", context.locale), context.report_meta.get("crew"))
    defect_text = f"{defect_count} defects" if context.locale == "en-US" else f"{defect_count} 处"
    _labeled(context, _label("defects_found", context.locale), defect_text)


def _write_results(
    context: ReportDocumentContext,
    overall_summary: str,
    sections: list[dict[str, Any]],
) -> None:
    add_heading(context.doc, _label("results", context.locale))
    add_subheading(context.doc, _label("pavement_condition", context.locale))
    _labeled(context, _label("overall", context.locale), overall_summary)
    for index, section in enumerate(sections, start=1):
        add_subheading(context.doc, _defect_heading(index, section, context.locale))
        _labeled(context, _label("features", context.locale), section.get("features"))
        if section.get("summary"):
            _labeled(context, _label("risk", context.locale), section.get("summary"))
        _write_defect_advice(context, section)


def _write_defect_advice(context: ReportDocumentContext, section: dict[str, Any]) -> None:
    _labeled(
        context,
        _label("recommendation", context.locale),
        section.get("recommendation") or section.get("chosen_method"),
    )
    _labeled(
        context,
        _label("process", context.locale),
        _method_text(section, context.locale),
    )
    _labeled(context, _label("quality", context.locale), _quality_text(context.locale))
    _labeled(context, _label("citations", context.locale), section.get("citations"))


def _write_recommendations(
    context: ReportDocumentContext,
    sections: list[dict[str, Any]],
    cost_summary: dict[str, Any] | None,
) -> None:
    add_heading(context.doc, _label("actions", context.locale))
    defect_text = (
        f"{len(sections)} defects; see pavement-condition entries above."
        if context.locale == "en-US"
        else f"{len(sections)} 处，详见路基路面状况各条。"
    )
    _labeled(context, _label("maintenance_needed", context.locale), defect_text)
    if cost_summary and cost_summary.get("total_cny") is not None:
        _labeled(context, _label("cost", context.locale), _cost_text(cost_summary, context.locale))
    else:
        _labeled(context, _label("cost", context.locale), _label("cost_sheet_ref", context.locale))
    _labeled(context, _label("coordination", context.locale), None)


def _write_attachments(context: ReportDocumentContext) -> None:
    add_heading(context.doc, _label("attachments", context.locale))
    _labeled(context, "1", _label("attachment_records", context.locale))
    _labeled(context, "2", _label("attachment_photos", context.locale))
    _labeled(context, "3", _label("attachment_cost", context.locale))


def _labeled(
    context: ReportDocumentContext,
    label: str,
    value: Any,
    *,
    placeholder: str | None = None,
) -> None:
    add_labeled_paragraph(
        context.doc,
        label,
        value,
        placeholder=placeholder,
        locale=context.locale,
    )


def _method_text(section: dict[str, Any], locale: str) -> str | None:
    method = section.get("chosen_method")
    if not method:
        return None
    if locale == "en-US":
        return (
            f"Use the confirmed method ({method}) and stage work according to site traffic control."
        )
    return f"建议采用{method}工艺，并按现场交通组织要求分段实施。"


def _cost_text(cost_summary: dict[str, Any], locale: str) -> str:
    defect = cost_summary.get("defect_subtotal_cny", 0)
    mobilization = cost_summary.get("mobilization_subtotal_cny", 0)
    total = cost_summary.get("total_cny", 0)
    if locale == "en-US":
        return (
            f"Defect treatment subtotal: {defect} CNY; mobilization: {mobilization} CNY; "
            f"estimated total: {total} CNY."
        )
    return f"病害处治小计 {defect} 元，进场费 {mobilization} 元，合计约 {total} 元。"


def _report_date(locale: str) -> str:
    today = date.today()
    if locale == "en-US":
        return today.isoformat()
    return f"{today.year}年{today.month:02d}月{today.day:02d}日"


def _defect_heading(index: int, section: dict[str, Any], locale: str) -> str:
    category = section.get("defect_category") or ("Defect" if locale == "en-US" else "病害")
    if locale == "en-US":
        return f"{index}. {section['location']} - {category}"
    return f"（{index}）{section['location']}{category}"


def _quality_text(locale: str) -> str:
    if locale == "en-US":
        return (
            "After construction, inspect against the applicable acceptance criteria and "
            "confirm the repaired area is compact, smooth, and free of contamination."
        )
    return "施工完成后按相应验收标准进行检查，确认处治部位密实、平整、无污染。"


def _label(key: str, locale: str) -> str:
    return report_label(key, locale)
