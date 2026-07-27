"""Render a formal road maintenance construction plan to DOCX."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from road_distress_agent.delivery.docx_format import (
    TITLE_TEXT,
    add_formal_paragraph,
    add_heading,
    add_labeled_paragraph,
    add_subheading,
    add_text,
    configure_formal_document,
)


@dataclass(frozen=True)
class PlanDocumentContext:
    doc: Any
    project: dict[str, Any]
    plan: dict[str, Any]
    locale: str


def write_construction_plan_document(
    *,
    project: dict[str, Any],
    plan: dict[str, Any],
    path: Path,
    locale: str = "zh-CN",
) -> Path:
    doc = Document()
    configure_formal_document(doc)
    context = PlanDocumentContext(doc=doc, project=project, plan=plan, locale=locale)
    _write_title(context)
    _write_overview(context)
    _write_defects(context)
    _write_schedule(context)
    _write_items(context, "materials")
    _write_items(context, "acceptance")
    doc.save(path)
    return path


def _write_title(context: PlanDocumentContext) -> None:
    paragraph = add_formal_paragraph(context.doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(paragraph, _label("title", context.locale), style=TITLE_TEXT)
    paragraph = add_formal_paragraph(context.doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(paragraph, f"{_label('plan_date', context.locale)}: {_today(context.locale)}")
    if context.plan.get("subject"):
        add_labeled_paragraph(
            context.doc,
            _label("subject", context.locale),
            context.plan.get("subject"),
            locale=context.locale,
        )


def _write_overview(context: PlanDocumentContext) -> None:
    add_heading(context.doc, _label("overview", context.locale))
    add_labeled_paragraph(
        context.doc,
        _label("project", context.locale),
        context.project.get("name"),
        locale=context.locale,
    )
    add_labeled_paragraph(
        context.doc,
        _label("segment", context.locale),
        context.project.get("segment"),
        locale=context.locale,
    )
    add_labeled_paragraph(
        context.doc,
        _label("crew", context.locale),
        context.project.get("crew"),
        locale=context.locale,
    )
    defects = context.plan.get("defects") or []
    count = f"{len(defects)} defects" if context.locale == "en-US" else f"{len(defects)} 处"
    add_labeled_paragraph(
        context.doc,
        _label("defect_count", context.locale),
        count,
        locale=context.locale,
    )


def _write_defects(context: PlanDocumentContext) -> None:
    defects = context.plan.get("defects") or []
    add_heading(context.doc, _label("defects", context.locale))
    if not defects:
        add_labeled_paragraph(
            context.doc,
            _label("defect_item", context.locale),
            None,
            locale=context.locale,
        )
        return
    for index, defect in enumerate(defects, start=1):
        add_subheading(context.doc, _defect_heading(index, defect, context.locale))
        add_labeled_paragraph(
            context.doc,
            _label("method", context.locale),
            defect.get("chosen_method"),
            locale=context.locale,
        )


def _write_schedule(context: PlanDocumentContext) -> None:
    schedule = context.plan.get("schedule") or {}
    add_heading(context.doc, _label("arrangement", context.locale))
    add_labeled_paragraph(
        context.doc,
        _label("body", context.locale),
        context.plan.get("body"),
        locale=context.locale,
    )
    add_labeled_paragraph(
        context.doc,
        _label("start_date", context.locale),
        schedule.get("start_date"),
        locale=context.locale,
    )
    add_labeled_paragraph(
        context.doc,
        _label("end_date", context.locale),
        schedule.get("end_date"),
        locale=context.locale,
    )
    add_labeled_paragraph(
        context.doc,
        _label("total_days", context.locale),
        schedule.get("total_days"),
        locale=context.locale,
    )


def _write_items(context: PlanDocumentContext, key: str) -> None:
    items = context.plan.get(key) or []
    add_heading(context.doc, _label(key, context.locale))
    if not items:
        add_labeled_paragraph(
            context.doc,
            _label("item", context.locale),
            None,
            locale=context.locale,
        )
        return
    for index, item in enumerate(items, start=1):
        add_labeled_paragraph(context.doc, str(index), item, locale=context.locale)


def _defect_heading(index: int, defect: dict[str, Any], locale: str) -> str:
    location = defect.get("location")
    category = defect.get("defect_category") or ("Defect" if locale == "en-US" else "病害")
    if locale == "en-US":
        return f"{index}. {location or '[Please provide Location]'} - {category}"
    return f"（{index}）{location or '[请补充位置]'}{category}"


def _today(locale: str) -> str:
    today = date.today()
    if locale == "en-US":
        return today.isoformat()
    return f"{today.year}年{today.month:02d}月{today.day:02d}日"


def _label(key: str, locale: str) -> str:
    labels = _LABELS_EN if locale == "en-US" else _LABELS_ZH
    return labels[key]


_LABELS_ZH = {
    "title": "道路维护施工方案",
    "plan_date": "编制日期",
    "subject": "方案名称",
    "overview": "一、工程概况",
    "project": "巡查任务",
    "segment": "施工路段",
    "crew": "责任班组",
    "defect_count": "处治病害",
    "defects": "二、病害与处治方案",
    "defect_item": "病害条目",
    "method": "处治方法",
    "arrangement": "三、施工组织安排",
    "body": "施工安排",
    "start_date": "计划开工日期",
    "end_date": "计划完工日期",
    "total_days": "计划工期（天）",
    "materials": "四、材料清单",
    "acceptance": "五、质量验收",
    "item": "条目",
}

_LABELS_EN = {
    "title": "Road Maintenance Construction Plan",
    "plan_date": "Plan Date",
    "subject": "Plan Name",
    "overview": "1. Project Overview",
    "project": "Inspection Task",
    "segment": "Work Segment",
    "crew": "Responsible Crew",
    "defect_count": "Defects to Treat",
    "defects": "2. Defects and Treatment Plan",
    "defect_item": "Defect Item",
    "method": "Treatment Method",
    "arrangement": "3. Construction Arrangement",
    "body": "Arrangement",
    "start_date": "Planned Start Date",
    "end_date": "Planned Completion Date",
    "total_days": "Planned Duration (days)",
    "materials": "4. Materials",
    "acceptance": "5. Quality Acceptance",
    "item": "Item",
}
