"""Formal Chinese document formatting helpers for delivery DOCX files."""

from __future__ import annotations

from dataclasses import dataclass
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

TITLE_FONT = "方正小标宋简体"
HEADING_FONT = "黑体"
SUBHEADING_FONT = "楷体_GB2312"
BODY_FONT = "仿宋_GB2312"
TITLE_SIZE = Pt(22)
BODY_SIZE = Pt(16)
LINE_SPACING_POINTS = 29
LINE_SPACING = Pt(LINE_SPACING_POINTS)
TWIPS_PER_POINT = 20
LINE_SPACING_TWIPS = LINE_SPACING_POINTS * TWIPS_PER_POINT
RED = RGBColor(192, 0, 0)
MARGINS_MM = {"top": 37, "bottom": 35, "left": 28, "right": 26}


@dataclass(frozen=True, kw_only=True)
class TextStyle:
    font: str
    size: Any
    bold: bool = False


TITLE_TEXT = TextStyle(font=TITLE_FONT, size=TITLE_SIZE, bold=True)
HEADING_TEXT = TextStyle(font=HEADING_FONT, size=BODY_SIZE, bold=True)
SUBHEADING_TEXT = TextStyle(font=SUBHEADING_FONT, size=BODY_SIZE, bold=True)
BODY_TEXT = TextStyle(font=BODY_FONT, size=BODY_SIZE)


def configure_formal_document(doc: Any) -> None:
    _configure_sections(doc)
    _set_style_font(doc.styles["Normal"], BODY_TEXT)


def add_formal_paragraph(
    doc: Any,
    *,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> Any:
    paragraph = doc.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = LINE_SPACING
    if alignment is not None:
        paragraph.alignment = alignment
    return paragraph


def add_text(
    paragraph: Any,
    text: str,
    *,
    style: TextStyle = BODY_TEXT,
    color: RGBColor | None = None,
) -> Any:
    run = paragraph.add_run(text)
    run.bold = style.bold
    run.font.size = style.size
    if color is not None:
        run.font.color.rgb = color
    _set_run_font(run, style.font)
    return run


def add_value_text(
    paragraph: Any,
    value: Any,
    placeholder: str,
    *,
    style: TextStyle = BODY_TEXT,
    locale: str = "zh-CN",
) -> None:
    text = str(value).strip() if value is not None else ""
    if text:
        add_text(paragraph, text, style=style)
        return
    missing = f"[Please provide {placeholder}]" if locale == "en-US" else f"[请补充{placeholder}]"
    add_text(paragraph, missing, style=style, color=RED)


def add_heading(doc: Any, text: str) -> None:
    paragraph = add_formal_paragraph(doc)
    add_text(paragraph, text, style=HEADING_TEXT)


def add_subheading(doc: Any, text: str) -> None:
    paragraph = add_formal_paragraph(doc)
    add_text(paragraph, text, style=SUBHEADING_TEXT)


def add_labeled_paragraph(
    doc: Any,
    label: str,
    value: Any,
    *,
    placeholder: str | None = None,
    locale: str = "zh-CN",
) -> None:
    paragraph = add_formal_paragraph(doc)
    add_text(paragraph, f"{label}: " if locale == "en-US" else f"{label}：")
    add_value_text(paragraph, value, placeholder or label, locale=locale)


def _configure_sections(doc: Any) -> None:
    for section in doc.sections:
        section.top_margin = Mm(MARGINS_MM["top"])
        section.bottom_margin = Mm(MARGINS_MM["bottom"])
        section.left_margin = Mm(MARGINS_MM["left"])
        section.right_margin = Mm(MARGINS_MM["right"])
        _set_doc_grid(section)


def _set_doc_grid(section: Any) -> None:
    section_properties = section._sectPr
    grid = section_properties.find(qn("w:docGrid"))
    if grid is None:
        grid = OxmlElement("w:docGrid")
        section_properties.append(grid)
    grid.set(qn("w:linePitch"), str(LINE_SPACING_TWIPS))


def _set_style_font(style: Any, text_style: TextStyle) -> None:
    style.font.name = text_style.font
    style.font.size = text_style.size
    rpr = style._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        style._element.append(rpr)
    _set_fonts(rpr, text_style.font)


def _set_run_font(run: Any, font_name: str) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    _set_fonts(rpr, font_name)


def _set_fonts(rpr: Any, font_name: str) -> None:
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), font_name)
